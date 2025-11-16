"""
Document Processing Service
Handles PDF, DOCX, CSV, and TXT file processing
"""

import io
import csv
from typing import List, Dict, Any, BinaryIO
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Process various document types for RAG
    """
    
    @staticmethod
    def process_csv(file_content: bytes) -> Dict[str, Any]:
        """
        Process CSV file and extract structured data
        """
        try:
            content = file_content.decode('utf-8')
            csv_reader = csv.DictReader(io.StringIO(content))
            
            rows = list(csv_reader)
            
            if not rows:
                return {"error": "Empty CSV file", "chunks": []}
            
            # Create text chunks from CSV rows
            chunks = []
            headers = list(rows[0].keys())
            
            # Create a summary chunk
            summary = f"CSV Document with {len(rows)} rows and {len(headers)} columns.\n"
            summary += f"Columns: {', '.join(headers)}\n\n"
            chunks.append({
                "text": summary,
                "type": "summary",
                "metadata": {
                    "row_count": len(rows),
                    "column_count": len(headers),
                    "columns": headers
                }
            })
            
            # Create chunks for each row (limit to first 100 rows for efficiency)
            for idx, row in enumerate(rows[:100]):
                row_text = f"Row {idx + 1}:\n"
                row_text += "\n".join([f"{k}: {v}" for k, v in row.items()])
                
                chunks.append({
                    "text": row_text,
                    "type": "data_row",
                    "metadata": {
                        "row_number": idx + 1,
                        "data": row
                    }
                })
            
            return {
                "chunks": chunks,
                "total_rows": len(rows),
                "columns": headers,
                "file_type": "csv"
            }
            
        except Exception as e:
            logger.error(f"Error processing CSV: {e}")
            return {"error": str(e), "chunks": []}
    
    @staticmethod
    def process_text(file_content: bytes) -> Dict[str, Any]:
        """
        Process plain text file
        """
        try:
            content = file_content.decode('utf-8')
            
            # Split into paragraphs
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            
            chunks = []
            for idx, para in enumerate(paragraphs):
                if len(para) > 50:  # Skip very short paragraphs
                    chunks.append({
                        "text": para,
                        "type": "paragraph",
                        "metadata": {
                            "paragraph_number": idx + 1,
                            "length": len(para)
                        }
                    })
            
            return {
                "chunks": chunks,
                "total_paragraphs": len(paragraphs),
                "file_type": "text"
            }
            
        except Exception as e:
            logger.error(f"Error processing text: {e}")
            return {"error": str(e), "chunks": []}
    
    @staticmethod
    def process_pdf(file_content: bytes) -> Dict[str, Any]:
        """
        Process PDF file (requires pypdf or similar)
        For now, returns placeholder
        """
        try:
            # Try to import pypdf
            try:
                from pypdf import PdfReader
                
                pdf_file = io.BytesIO(file_content)
                reader = PdfReader(pdf_file)
                
                chunks = []
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        chunks.append({
                            "text": text,
                            "type": "pdf_page",
                            "metadata": {
                                "page_number": page_num + 1
                            }
                        })
                
                return {
                    "chunks": chunks,
                    "total_pages": len(reader.pages),
                    "file_type": "pdf"
                }
                
            except ImportError:
                # Fallback when pypdf not available
                return {
                    "error": "PDF processing requires 'pypdf' library. Install with: pip install pypdf",
                    "chunks": [],
                    "file_type": "pdf"
                }
                
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return {"error": str(e), "chunks": []}
    
    @staticmethod
    def process_docx(file_content: bytes) -> Dict[str, Any]:
        """
        Process DOCX file (requires python-docx)
        """
        try:
            # Try to import python-docx
            try:
                from docx import Document
                
                docx_file = io.BytesIO(file_content)
                doc = Document(docx_file)
                
                chunks = []
                for para_num, para in enumerate(doc.paragraphs):
                    text = para.text.strip()
                    if text and len(text) > 50:
                        chunks.append({
                            "text": text,
                            "type": "docx_paragraph",
                            "metadata": {
                                "paragraph_number": para_num + 1
                            }
                        })
                
                return {
                    "chunks": chunks,
                    "total_paragraphs": len(doc.paragraphs),
                    "file_type": "docx"
                }
                
            except ImportError:
                return {
                    "error": "DOCX processing requires 'python-docx' library. Install with: pip install python-docx",
                    "chunks": [],
                    "file_type": "docx"
                }
                
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            return {"error": str(e), "chunks": []}
    
    @staticmethod
    def process_file(filename: str, file_content: bytes) -> Dict[str, Any]:
        """
        Process file based on extension
        """
        extension = filename.lower().split('.')[-1]
        
        processors = {
            'csv': DocumentProcessor.process_csv,
            'txt': DocumentProcessor.process_text,
            'pdf': DocumentProcessor.process_pdf,
            'docx': DocumentProcessor.process_docx,
            'doc': DocumentProcessor.process_docx,
        }
        
        processor = processors.get(extension)
        
        if not processor:
            return {
                "error": f"Unsupported file type: {extension}",
                "chunks": []
            }
        
        result = processor(file_content)
        result["filename"] = filename
        result["file_extension"] = extension
        
        return result


class SimpleVectorStore:
    """
    Simple in-memory vector store for RAG
    Uses basic cosine similarity without heavy dependencies
    """
    
    def __init__(self):
        self.documents = []
        self.embeddings = []
    
    def add_documents(self, chunks: List[Dict[str, Any]]):
        """
        Add document chunks to store
        For now, stores without embeddings (can add later with sentence-transformers)
        """
        for chunk in chunks:
            self.documents.append(chunk)
    
    def simple_search(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """
        Simple keyword-based search (no embeddings)
        """
        query_lower = query.lower()
        query_words = set(query_lower.split())
        
        # Score documents by word overlap
        scored_docs = []
        for doc in self.documents:
            text_lower = doc['text'].lower()
            text_words = set(text_lower.split())
            
            # Calculate overlap score
            overlap = len(query_words & text_words)
            if overlap > 0:
                scored_docs.append((overlap, doc))
        
        # Sort by score and return top k
        scored_docs.sort(reverse=True, key=lambda x: x[0])
        return [doc for score, doc in scored_docs[:top_k]]
    
    def get_all_documents(self) -> List[Dict[str, Any]]:
        """Return all stored documents"""
        return self.documents
    
    def clear(self):
        """Clear all documents"""
        self.documents = []
        self.embeddings = []
